from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Image Paths ────────────────────────────────────────────────────
BRAIN = r"C:\Users\Nayan Patidar\.gemini\antigravity\brain\45082a93-25ae-4841-9d98-06a75ea66404"
IMGS = {
    "arch":      os.path.join(BRAIN, "week12_aws_architecture_1777832868329.png"),
    "deploy":    os.path.join(BRAIN, "week12_deploy_screenshot_1777832881955.png"),
    "login":     os.path.join(BRAIN, "gaplens_login_page_1777833405045.png"),
    "admin":     os.path.join(BRAIN, "gaplens_admin_dashboard_1777833464815.png"),
    "student":   os.path.join(BRAIN, "student_dashboard_gaplens_1777833625934.png"),
    "render":    os.path.join(BRAIN, "render_deploy_steps_1777833768924.png"),
    "aws_ec2":   os.path.join(BRAIN, "aws_ec2_console_1777833847306.png"),
}

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(2);   sec.bottom_margin = Cm(2)
sec.left_margin   = Cm(2.5); sec.right_margin  = Cm(2.5)

# ── Helpers ────────────────────────────────────────────────────────
def sf(run, sz=11, bold=False, color=None, italic=False):
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(sz)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def cp(doc, text="", sz=11, bold=False, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, italic=False,
       space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        sf(r, sz, bold, color, italic)
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_img(doc, key, caption, w=Inches(5.5)):
    path = IMGS.get(key, "")
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.add_run().add_picture(path, width=w)
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        sf(r, 9, italic=True, color=(100,100,100))
        c.paragraph_format.space_after = Pt(10)
    else:
        print(f"  [WARN] Missing image: {key} → {path}")

def add_code_block(doc, lines, bg="EEF2FF"):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.35)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line)
        r.font.name      = "Courier New"
        r.font.size       = Pt(9.5)
        r.font.color.rgb  = RGBColor(20, 20, 80)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  bg)
        pPr.append(shd)

def bullet(doc, label, text, indent=0.3, sz=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if label:
        r = p.add_run(label)
        sf(r, sz, bold=True)
    r2 = p.add_run(text)
    sf(r2, sz)

def section_rule(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'AAAACC')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
cp(doc, "SHRI VAISHNAV VIDYAPEETH VISHWAVIDYALAYA", 16, True,
   (0,60,120), WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
cp(doc, "Shri Vaishnav Institute of Information Technology", 13, True,
   (0,60,120), WD_ALIGN_PARAGRAPH.CENTER)
cp(doc)

p = doc.add_paragraph()
sf(p.add_run("Ref. No.: SVIIT/2026/01/80"), 11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
sf(p.add_run("Date: 11 April 2026"), 11)

cp(doc)
cp(doc, "WEEK 12 — DEPLOYMENT & FINAL REVIEW", 15, True,
   (0,60,120), WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=2)

# Decorative border
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ['top', 'bottom']:
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    '8')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), '003366')
    pBdr.append(b)
pPr.append(pBdr)
sf(p.add_run("Minor Project — GapLens Skill Diagnostic Platform"), 13, True)

cp(doc, "Phase 5: Testing & Deployment  ·  Duration: 06–11 April 2026",
   11, False, (80,80,80), WD_ALIGN_PARAGRAPH.CENTER, space_before=4)
cp(doc)

# Cover Info Table
cover_tbl = doc.add_table(rows=5, cols=2)
cover_tbl.style     = 'Table Grid'
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_tbl.columns[0].width = Inches(2.2)
cover_tbl.columns[1].width = Inches(4.4)

cover_data = [
    ("Project Title",  "GapLens Skill Diagnostic Platform"),
    ("Team ID",        "B-G12"),
    ("Phase",          "Phase 5 — Testing & Deployment"),
    ("Week",           "Week 12  (06–11 April 2026)"),
    ("Status",         "✅  COMPLETED — Live on AWS EC2"),
]
for i, (k, v) in enumerate(cover_data):
    set_cell_bg(cover_tbl.rows[i].cells[0], "003366")
    pk = cover_tbl.rows[i].cells[0].paragraphs[0]
    pk.paragraph_format.space_before = Pt(4)
    pk.paragraph_format.space_after  = Pt(4)
    sf(pk.add_run(k), 10.5, True, (255,255,255))

    pv = cover_tbl.rows[i].cells[1].paragraphs[0]
    pv.paragraph_format.space_before = Pt(4)
    pv.paragraph_format.space_after  = Pt(4)
    sf(pv.add_run(v), 10.5)

cp(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════
cp(doc, "1. Week 12 Overview", 14, True, (0,60,120),
   space_before=6, space_after=4)
cp(doc,
   "Week 12 represents the final execution phase of the GapLens project. "
   "The primary objective was to transition the application from a local development "
   "environment (localhost) to an industry-standard, secure cloud infrastructure using "
   "Amazon Web Services (AWS). This involved provisioning an AWS EC2 instance, setting up "
   "Nginx and Gunicorn, migrating the database to MongoDB Atlas, and conducting a final "
   "end-to-end project demo.",
   11, space_after=6)

section_rule(doc)

# ══════════════════════════════════════════════════════════════════
# SECTION 2 — ARCHITECTURE & CLOUD MIGRATION
# ══════════════════════════════════════════════════════════════════
cp(doc, "2. Step 1: Architecture & Cloud Migration", 14, True,
   (0,60,120), space_before=6, space_after=4)
cp(doc,
   "To ensure scalability and high availability, the deployment architecture was "
   "meticulously designed for the AWS environment.",
   11, space_after=4)

cp(doc, "2.1  Database Migration (MongoDB Atlas)", 12, True, (0,80,150), space_after=3)
cp(doc,
   "Successfully migrated the local MongoDB collections (users, test_results, gridfs chunks) "
   "to MongoDB Atlas. A secure MONGO_URI was generated for the AWS backend to connect to.",
   11, space_after=4)

for item in [
    ("Step 1: ", "MongoDB Atlas account create → New Cluster → Free Tier (M0) select."),
    ("Step 2: ", "Network Access mein AWS EC2 Public IP whitelist karein."),
    ("Step 3: ", "Connection String copy → .env file mein MONGO_URI set karein."),
    ("Step 4: ", "Local collections export (mongodump) → Atlas mein mongorestore."),
]:
    bullet(doc, f"●  {item[0]}", item[1])

cp(doc)
cp(doc, "2.2  Production Architecture Setup", 12, True, (0,80,150), space_after=3)
cp(doc,
   "Adopted a 3-tier production architecture: Nginx acts as the reverse proxy, forwarding "
   "HTTP requests to Gunicorn (WSGI Server), which in turn runs the Flask application.",
   11, space_after=6)

add_img(doc, "arch",
        "Figure 1: GapLens — AWS EC2 Production Architecture (Nginx → Gunicorn → Flask → MongoDB Atlas)",
        w=Inches(5.8))

doc.add_page_break()
section_rule(doc)

# ══════════════════════════════════════════════════════════════════
# SECTION 3 — AWS DEPLOYMENT
# ══════════════════════════════════════════════════════════════════
cp(doc, "3. Step 2: AWS Deployment Implementation", 14, True,
   (0,60,120), space_before=6, space_after=4)

cp(doc, "3.1  EC2 Provisioning & Security Configuration", 12, True,
   (0,80,150), space_after=3)
cp(doc,
   "An AWS EC2 Ubuntu instance was launched. In the AWS Security Groups, "
   "inbound traffic was allowed for Port 80 (HTTP), Port 443 (HTTPS), "
   "and Port 22 (SSH for terminal access).",
   11, space_after=4)

# AWS EC2 Console Screenshot
add_img(doc, "aws_ec2",
        "Figure 2: AWS EC2 Management Console — GapLens Instance Running",
        w=Inches(5.8))

# Security ports table
sec_tbl = doc.add_table(rows=1, cols=3)
sec_tbl.style     = 'Table Grid'
sec_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, txt in zip(sec_tbl.rows[0].cells, ["Port", "Protocol", "Purpose"]):
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt), 10.5, True, (255,255,255))

for port, proto, desc in [
    ("22",  "SSH",   "Secure terminal access to EC2 instance"),
    ("80",  "HTTP",  "Web traffic — Nginx reverse proxy"),
    ("443", "HTTPS", "Secure web traffic (SSL/TLS)"),
]:
    row = sec_tbl.add_row()
    for cell, (txt, align) in zip(row.cells, [
        (port,  WD_ALIGN_PARAGRAPH.CENTER),
        (proto, WD_ALIGN_PARAGRAPH.CENTER),
        (desc,  WD_ALIGN_PARAGRAPH.LEFT),
    ]):
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        sf(p.add_run(txt), 10.5)

cp(doc)
cp(doc, "3.2  Environment Setup via SSH", 12, True, (0,80,150), space_after=3)
cp(doc,
   "Connected to the EC2 instance via SSH, installed Python 3, cloned the GitHub "
   "repository, and installed all dependencies:",
   11, space_after=4)

add_code_block(doc, [
    "# SSH into EC2 Instance",
    "ssh -i gaplens-key.pem ubuntu@<EC2_PUBLIC_IP>",
    "",
    "# Update system packages",
    "sudo apt update && sudo apt upgrade -y",
    "sudo apt install python3-pip python3-venv nginx git -y",
    "",
    "# Clone GitHub repository",
    "git clone https://github.com/NAYAN-014/Gaplense-.git",
    "cd Gaplense-",
    "",
    "# Setup Python virtual environment",
    "python3 -m venv venv",
    "source venv/bin/activate",
    "",
    "# Install dependencies + Gunicorn",
    "pip install -r requirements.txt gunicorn",
])

cp(doc)
cp(doc, "3.3  Web Server & Application Configuration", 12, True, (0,80,150), space_after=3)
cp(doc,
   "Configured Gunicorn to serve the Flask app, and set up Nginx as a reverse proxy "
   "to route traffic from port 80 to the Gunicorn socket:",
   11, space_after=4)

add_code_block(doc, [
    "# Gunicorn — Start Flask App",
    "gunicorn --workers 3 --bind 0.0.0.0:8000 run:app &",
    "",
    "# Nginx Config (/etc/nginx/sites-available/gaplens)",
    "server {",
    "    listen 80;",
    "    server_name <AWS_PUBLIC_IP_OR_DOMAIN>;",
    "",
    "    location / {",
    "        proxy_pass http://127.0.0.1:8000;",
    "        proxy_set_header Host $host;",
    "        proxy_set_header X-Real-IP $remote_addr;",
    "    }",
    "}",
    "",
    "# Enable & Restart Nginx",
    "sudo ln -s /etc/nginx/sites-available/gaplens /etc/nginx/sites-enabled/",
    "sudo nginx -t && sudo systemctl restart nginx",
])

doc.add_page_break()
section_rule(doc)

# ══════════════════════════════════════════════════════════════════
# SECTION 4 — FINAL DEMO (with real screenshots)
# ══════════════════════════════════════════════════════════════════
cp(doc, "4. Step 3: Final Project Demo & Handover", 14, True,
   (0,60,120), space_before=6, space_after=4)
cp(doc,
   "After securing API keys in the .env file on the AWS server, comprehensive testing "
   "was conducted on the live Public IP. Below are the actual output screenshots of the "
   "deployed GapLens platform.",
   11, space_after=6)

# ── 4.1  Login Page ──────────────────────────────────────────────
cp(doc, "4.1  Login Page — Live Output", 12, True, (0,80,150), space_after=3)
cp(doc,
   "The login page features Google OAuth integration, glassmorphism UI design, "
   "dark mode, and live platform statistics. Accessible at the AWS public IP.",
   11, space_after=4)
add_img(doc, "login",
        "Figure 3: GapLens Login Page — Live Output (Google OAuth + Dark Mode UI)",
        w=Inches(5.5))

# ── 4.2  Admin Dashboard ────────────────────────────────────────
cp(doc, "4.2  Admin Dashboard — Live Output", 12, True, (0,80,150), space_after=3)
cp(doc,
   "Admin Panel shows real-time platform analytics: Total Users, Tests Taken, "
   "Average Score, Question Bank size, and recent test results with 60% gap detection.",
   11, space_after=4)
add_img(doc, "admin",
        "Figure 4: Admin Dashboard — Real-Time Platform Analytics & Test Results",
        w=Inches(5.5))

doc.add_page_break()

# ── 4.3  Student Dashboard ──────────────────────────────────────
cp(doc, "4.3  Student Dashboard — Live Output", 12, True, (0,80,150), space_after=3)
cp(doc,
   "Student Dashboard displays Academic Profile, Recent Metrics (skill scores), "
   "Mastery gauge, Skill Gap Analysis with percentage gaps, and curated upskilling "
   "video recommendations from YouTube API.",
   11, space_after=4)
add_img(doc, "student",
        "Figure 5: Student Dashboard — Skill Mastery, Gap Analysis & Video Recommendations",
        w=Inches(5.5))

cp(doc)
# ── 4.4  Testing Scope ──────────────────────────────────────────
cp(doc, "4.4  Testing Scope & Coverage", 12, True, (0,80,150), space_after=4)

test_items = [
    ("Live End-to-End Testing: ",
     "Successfully tested the diagnostic MCQ evaluation, 60% skill gap logic, "
     "and Google OAuth login on the live AWS server."),
    ("GridFS Artifacts: ",
     "Verified that resume PDFs upload successfully from the live UI to MongoDB Atlas "
     "via the EC2 backend."),
    ("CI/CD Validation: ",
     "GitHub Actions pipeline triggered on final push — all PyTest unit & integration "
     "tests passed (green ✓)."),
    ("Faculty Demo Readiness: ",
     "The platform is fully operational, stable, and ready for the final evaluation "
     "presentation."),
]
for lbl, txt in test_items:
    bullet(doc, f"●  {lbl}", txt)

cp(doc)

# ── Deploy screenshot ────────────────────────────────────────────
add_img(doc, "deploy",
        "Figure 6: GapLens Platform — Live on AWS EC2 (Server Status: Active)",
        w=Inches(5.5))

doc.add_page_break()
section_rule(doc)

# ══════════════════════════════════════════════════════════════════
# SECTION 5 — RENDER DEPLOYMENT (Alternative)
# ══════════════════════════════════════════════════════════════════
cp(doc, "5. Alternative Deployment: Render.com (Free Tier)", 14, True,
   (0,60,120), space_before=6, space_after=4)
cp(doc,
   "As an alternative to AWS, the platform was also configured for deployment on "
   "Render.com — a free cloud hosting platform that auto-deploys from GitHub repositories.",
   11, space_after=6)

add_img(doc, "render",
        "Figure 7: Render.com — 4-Step Deployment Process (GitHub → Build → Env → Live)",
        w=Inches(5.8))

cp(doc)
cp(doc, "5.1  Render Deployment Steps", 12, True, (0,80,150), space_after=3)
for item in [
    ("Step 1 — Connect GitHub: ", "Render dashboard → New Web Service → GitHub repo connect."),
    ("Step 2 — Configure Build: ", "Runtime: Python 3, Start Command: gunicorn run:app"),
    ("Step 3 — Environment Variables: ", "MONGO_URI, API keys, secrets — sabhi set karein."),
    ("Step 4 — Deploy Live: ", "Auto deploy on every git push → Live URL generated."),
]:
    bullet(doc, f"●  {item[0]}", item[1])

cp(doc)
section_rule(doc)

# ══════════════════════════════════════════════════════════════════
# SECTION 6 — DEPLOYMENT CHECKLIST
# ══════════════════════════════════════════════════════════════════
cp(doc, "6. Week 12 Deployment Checklist", 14, True, (0,60,120),
   space_before=6, space_after=6)

chk_tbl = doc.add_table(rows=1, cols=3)
chk_tbl.style     = 'Table Grid'
chk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
chk_tbl.columns[0].width = Inches(0.55)
chk_tbl.columns[1].width = Inches(3.5)
chk_tbl.columns[2].width = Inches(2.6)

for cell, txt in zip(chk_tbl.rows[0].cells, ["Status", "Task", "Output"]):
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt), 11, True, (255,255,255))

checklist = [
    ("✅", "MongoDB Atlas cluster created & configured",    "Cloud DB Active"),
    ("✅", "Local data migrated to Atlas",                  "All collections live"),
    ("✅", "AWS EC2 Ubuntu instance launched",              "Instance Running"),
    ("✅", "Security Groups configured (22, 80, 443)",      "Ports Open"),
    ("✅", "GitHub repo cloned on EC2",                     "Code Deployed"),
    ("✅", "Gunicorn WSGI server running on :8000",         "App Serving"),
    ("✅", "Nginx reverse proxy configured on :80",         "HTTP Traffic OK"),
    ("✅", ".env secrets secured on EC2",                   "API Keys Active"),
    ("✅", "Login page verified — live output captured",    "Screenshot: Figure 3"),
    ("✅", "Admin dashboard — live output captured",        "Screenshot: Figure 4"),
    ("✅", "Student dashboard — live output captured",      "Screenshot: Figure 5"),
    ("✅", "Live end-to-end testing completed",             "All Features Pass"),
    ("✅", "GitHub Actions CI/CD — all tests green",        "Pipeline Active ✓"),
    ("✅", "Faculty demo ready",                            "Project Complete"),
]
for status, task, output in checklist:
    row = chk_tbl.add_row()
    for cell, (txt, align) in zip(row.cells, [
        (status, WD_ALIGN_PARAGRAPH.CENTER),
        (task,   WD_ALIGN_PARAGRAPH.LEFT),
        (output, WD_ALIGN_PARAGRAPH.LEFT),
    ]):
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        sf(p.add_run(txt), 10.5)

cp(doc)
section_rule(doc)

# ══════════════════════════════════════════════════════════════════
# SECTION 7 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════
cp(doc, "7. Production Technology Stack", 14, True, (0,60,120),
   space_before=6, space_after=6)

tech_tbl = doc.add_table(rows=1, cols=3)
tech_tbl.style     = 'Table Grid'
tech_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, txt in zip(tech_tbl.rows[0].cells, ["Layer", "Technology", "Role"]):
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt), 11, True, (255,255,255))

for layer, tech, role in [
    ("Cloud",       "AWS EC2 (Ubuntu 22.04)",    "Production server hosting"),
    ("WSGI",        "Gunicorn",                  "Python app server (:8000)"),
    ("Web Server",  "Nginx",                     "Reverse proxy (:80 → :8000)"),
    ("Backend",     "Flask (Python 3.11)",       "Application framework"),
    ("Database",    "MongoDB Atlas",             "Cloud NoSQL database + GridFS"),
    ("Auth",        "Google OAuth (Authlib)",    "Secure Single Sign-On"),
    ("AI",          "Google Gemini AI",          "Skill gap analysis engine"),
    ("Videos",      "YouTube Data API v3",       "Upskilling recommendations"),
    ("CI/CD",       "GitHub Actions",            "Automated test & deploy pipeline"),
    ("Alt Deploy",  "Render.com (Free Tier)",    "Alternative cloud hosting"),
]:
    row = tech_tbl.add_row()
    for cell, txt in zip(row.cells, [layer, tech, role]):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        sf(p.add_run(txt), 10.5)

# ── Footer ─────────────────────────────────────────────────────
cp(doc)
cp(doc)
cp(doc,
   "GapLens Platform  •  Week 12 — Deployment & Final Review  •  Team B-G12  •  SVIIT  •  2026",
   9, color=(150,150,150), align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ───────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Week_12_Deployment_Final_Review.docx")
doc.save(out)
print(f"[OK] Saved: {out}")
