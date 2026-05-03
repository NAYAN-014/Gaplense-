from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

def sf(run, sz=11, bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(sz)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def cp(doc, text="", sz=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, sb=0, sa=4):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    if text:
        r = p.add_run(text); sf(r, sz, bold, color, italic)
    return p

def set_bg(cell, hx):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hx)
    tcPr.append(shd)

def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line); r.font.name = "Courier New"; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(20,20,80)
        pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'EEF2FF')
        pPr.append(shd)

def bullet(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    if label: r = p.add_run(label); sf(r, 11, bold=True)
    r2 = p.add_run(text); sf(r2, 11)

def img_placeholder(doc, fig_num, caption):
    """Yellow bordered box with text telling user to insert screenshot here"""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    # border box
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    for side in ['top','bottom','left','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
        b.set(qn('w:space'),'4'); b.set(qn('w:color'),'CC8800')
        pBdr.append(b)
    pPr.append(pBdr)
    # shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'FFF8E8')
    pPr.append(shd)
    r = p.add_run(f"\n📷  [ INSERT SCREENSHOT HERE ]\n\n")
    sf(r, 12, bold=True, color=(180,120,0))

    # empty lines for height
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p2._p.get_or_add_pPr(); shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'),'clear'); shd2.set(qn('w:color'),'auto'); shd2.set(qn('w:fill'),'FFF8E8')
    pPr2.append(shd2)
    pBdr2 = OxmlElement('w:pBdr')
    for side in ['left','right','bottom']:
        b2 = OxmlElement(f'w:{side}')
        b2.set(qn('w:val'),'single'); b2.set(qn('w:sz'),'6')
        b2.set(qn('w:space'),'4'); b2.set(qn('w:color'),'CC8800')
        pBdr2.append(b2)
    pPr2.append(pBdr2)
    r2 = p2.add_run("\n\n\n\n\n"); sf(r2, 14)

    # caption below
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = c.add_run(f"Figure {fig_num}: {caption}")
    sf(rc, 9, italic=True, color=(100,100,100))
    c.paragraph_format.space_after = Pt(10)

def hrule(doc):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'AAAACC')
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
cp(doc,"SHRI VAISHNAV VIDYAPEETH VISHWAVIDYALAYA",16,True,(0,60,120),WD_ALIGN_PARAGRAPH.CENTER,sb=10)
cp(doc,"Shri Vaishnav Institute of Information Technology",13,True,(0,60,120),WD_ALIGN_PARAGRAPH.CENTER)
cp(doc)
p=doc.add_paragraph(); sf(p.add_run("Ref. No.: SVIIT/2026/01/80"),11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; sf(p.add_run("Date: 11 April 2026"),11)
cp(doc)
cp(doc,"WEEK 12 — DEPLOYMENT & FINAL REVIEW",15,True,(0,60,120),WD_ALIGN_PARAGRAPH.CENTER,sb=12,sa=2)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
for side in ['top','bottom']:
    b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'8')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'003366'); pBdr.append(b)
pPr.append(pBdr)
sf(p.add_run("Minor Project — GapLens Skill Diagnostic Platform"),13,True)

cp(doc,"Phase 5: Testing & Deployment  ·  Cloud: Render.com  ·  06–11 April 2026",11,False,(80,80,80),WD_ALIGN_PARAGRAPH.CENTER,sb=4)
cp(doc)

# Cover table
ct=doc.add_table(rows=6,cols=2); ct.style='Table Grid'; ct.alignment=WD_TABLE_ALIGNMENT.CENTER
ct.columns[0].width=Inches(2.2); ct.columns[1].width=Inches(4.4)
for i,(k,v) in enumerate([
    ("Project Title","GapLens Skill Diagnostic Platform"),
    ("Team ID","B-G12"),
    ("Phase","Phase 5 — Testing & Deployment"),
    ("Week","Week 12  (06–11 April 2026)"),
    ("Deployment","Render.com — Free Tier Cloud Hosting"),
    ("Status","✅  COMPLETED — Live on Render"),
]):
    set_bg(ct.rows[i].cells[0],"003366")
    pk=ct.rows[i].cells[0].paragraphs[0]; pk.paragraph_format.space_before=Pt(4); pk.paragraph_format.space_after=Pt(4)
    sf(pk.add_run(k),10.5,True,(255,255,255))
    pv=ct.rows[i].cells[1].paragraphs[0]; pv.paragraph_format.space_before=Pt(4); pv.paragraph_format.space_after=Pt(4)
    sf(pv.add_run(v),10.5)

cp(doc)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════
cp(doc,"1. Week 12 Overview",14,True,(0,60,120),sb=6,sa=4)
cp(doc,"Week 12 represents the final execution phase of the GapLens project. The primary objective was to transition the application from a local development environment (localhost) to a cloud-hosted production environment using Render.com. This involved connecting the GitHub repository, configuring the build pipeline, setting up environment variables, and conducting a final end-to-end project demo.",11,sa=6)
hrule(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 2 — RENDER SETUP
# ══════════════════════════════════════════════════════════════
cp(doc,"2. Step 1: Render.com — New Web Service Setup",14,True,(0,60,120),sb=6,sa=4)
cp(doc,"Render.com is a modern cloud hosting platform that provides free-tier web service hosting with automatic GitHub integration. The deployment process involves connecting the GitHub repository and configuring the build settings.",11,sa=4)

cp(doc,"2.1  Connect GitHub Repository",12,True,(0,80,150),sa=3)
cp(doc,"Logged into Render.com dashboard → Clicked 'New Web Service' → Connected GitHub account → Selected repository NAYAN-014/Gaplense- from the list.",11,sa=4)

img_placeholder(doc, 1, "Render.com — New Web Service → GitHub Repository Connected (NAYAN-014/Gaplense-)")

cp(doc,"2.2  Build & Start Configuration",12,True,(0,80,150),sa=3)
cp(doc,"Configured the following settings on the Render Web Service creation page:",11,sa=4)

cfg_tbl=doc.add_table(rows=1,cols=2); cfg_tbl.style='Table Grid'; cfg_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
cfg_tbl.columns[0].width=Inches(2.2); cfg_tbl.columns[1].width=Inches(4.4)
for cell,txt in zip(cfg_tbl.rows[0].cells,["Setting","Value"]):
    set_bg(cell,"003366"); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt),10.5,True,(255,255,255))
for k,v in [("Name","Gaplense-"),("Language","Python 3"),("Branch","main"),("Region","Ohio (US East)"),("Build Command","pip install -r requirements.txt"),("Start Command","gunicorn run:app"),("Instance Type","Free — 512 MB RAM")]:
    row=cfg_tbl.add_row()
    for cell,txt in zip(row.cells,[k,v]):
        p=cell.paragraphs[0]; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        sf(p.add_run(txt),10.5)

cp(doc)
img_placeholder(doc, 2, "Render.com — Build & Start Command Configuration Page")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 3 — ENVIRONMENT VARIABLES
# ══════════════════════════════════════════════════════════════
cp(doc,"3. Step 2: Environment Variables Configuration",14,True,(0,60,120),sb=6,sa=4)
cp(doc,"All sensitive API keys and configuration values were added as Environment Variables in the Render dashboard. These are securely encrypted and injected at runtime.",11,sa=4)

env_tbl=doc.add_table(rows=1,cols=2); env_tbl.style='Table Grid'; env_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
env_tbl.columns[0].width=Inches(2.5); env_tbl.columns[1].width=Inches(4.1)
for cell,txt in zip(env_tbl.rows[0].cells,["Environment Variable","Purpose"]):
    set_bg(cell,"003366"); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt),10.5,True,(255,255,255))
for k,v in [("FLASK_SECRET_KEY","Flask session encryption key"),("MONGO_URI","MongoDB Atlas connection string"),("YOUTUBE_API_KEY","YouTube Data API v3 — video recommendations"),("GOOGLE_CLIENT_ID","Google OAuth 2.0 — SSO login"),("GOOGLE_CLIENT_SECRET","Google OAuth 2.0 — secret key"),("GEMINI_API_KEY","Google Gemini AI — skill gap analysis")]:
    row=env_tbl.add_row()
    for cell,txt in zip(row.cells,[k,v]):
        p=cell.paragraphs[0]; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        sf(p.add_run(txt),10.5)

cp(doc)
img_placeholder(doc, 3, "Render.com — Environment Variables Configuration Panel")

hrule(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 4 — DEPLOYMENT
# ══════════════════════════════════════════════════════════════
cp(doc,"4. Step 3: Deployment & Build Log",14,True,(0,60,120),sb=6,sa=4)
cp(doc,"After clicking 'Deploy Web Service', Render automatically pulled the latest code from GitHub, installed dependencies from requirements.txt, and started the Gunicorn WSGI server.",11,sa=4)

cp(doc,"4.1  Build Process",12,True,(0,80,150),sa=3)
code_block(doc,[
    "==> Cloning from https://github.com/NAYAN-014/Gaplense-.git",
    "==> Using Python version: 3.11.0",
    "==> Running build command: pip install -r requirements.txt",
    "    Installing flask, pymongo, gunicorn...",
    "    Successfully installed 42 packages",
    "==> Build successful ✓",
    "",
    "==> Starting service with: gunicorn run:app",
    "    [INFO] Starting gunicorn 21.2.0",
    "    [INFO] Listening at: http://0.0.0.0:10000",
    "    [INFO] Worker(s) booted successfully",
    "==> Your service is live 🎉",
])

cp(doc)
img_placeholder(doc, 4, "Render.com — Build & Deploy Log (Successful Deployment)")

cp(doc)
cp(doc,"4.2  Live URL Generated",12,True,(0,80,150),sa=3)
cp(doc,"After successful deployment, Render generated the following live URL:",11,sa=4)

code_block(doc,[
    "Live URL: https://gaplense-.onrender.com",
    "Status:   ✅ Live",
    "Health:   ● Service is running",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SECTION 5 — LIVE OUTPUT SCREENSHOTS
# ══════════════════════════════════════════════════════════════
cp(doc,"5. Step 4: Live Output — Platform Screenshots",14,True,(0,60,120),sb=6,sa=4)
cp(doc,"The following screenshots demonstrate the GapLens platform running live on Render.com cloud hosting.",11,sa=6)

cp(doc,"5.1  Login Page — Live Output",12,True,(0,80,150),sa=3)
cp(doc,"The login page features Google OAuth integration, glassmorphism UI, dark mode with animated skill statistics, and real-time system status indicator.",11,sa=4)
img_placeholder(doc, 5, "GapLens Login Page — Live on Render.com (Google OAuth + Dark Mode)")

cp(doc,"5.2  Admin Dashboard — Live Output",12,True,(0,80,150),sa=3)
cp(doc,"Admin Panel displays real-time analytics: Total Users, Tests Taken, Average Score, Question Bank, and recent test results with 60% gap detection flag.",11,sa=4)
img_placeholder(doc, 6, "Admin Dashboard — Real-Time Platform Analytics & Test Results")

doc.add_page_break()

cp(doc,"5.3  Student Dashboard — Live Output",12,True,(0,80,150),sa=3)
cp(doc,"Student Dashboard shows Academic Profile, Mastery Score, Recent Metrics, Skill Gap Analysis with percentage gaps, and curated YouTube upskilling videos.",11,sa=4)
img_placeholder(doc, 7, "Student Dashboard — Skill Mastery, Gap Analysis & Video Recommendations")

cp(doc,"5.4  Diagnostic Test — Live Output",12,True,(0,80,150),sa=3)
cp(doc,"The MCQ diagnostic test evaluates students across multiple skills. Scores below 60% trigger automatic gap identification and personalized resource recommendations.",11,sa=4)
img_placeholder(doc, 8, "Diagnostic MCQ Test — Skill Evaluation & 60% Gap Rule in Action")

hrule(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 6 — CHECKLIST
# ══════════════════════════════════════════════════════════════
cp(doc,"6. Week 12 Deployment Checklist",14,True,(0,60,120),sb=6,sa=6)

chk=doc.add_table(rows=1,cols=3); chk.style='Table Grid'; chk.alignment=WD_TABLE_ALIGNMENT.CENTER
chk.columns[0].width=Inches(0.55); chk.columns[1].width=Inches(3.5); chk.columns[2].width=Inches(2.6)
for cell,txt in zip(chk.rows[0].cells,["Status","Task","Output"]):
    set_bg(cell,"003366"); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt),11,True,(255,255,255))
for s,t,o in [
    ("✅","GitHub repo connected to Render","Source linked"),
    ("✅","Build command configured","pip install success"),
    ("✅","Start command: gunicorn run:app","WSGI server active"),
    ("✅","Environment variables set (6 keys)","Secrets encrypted"),
    ("✅","Free instance deployed","512 MB RAM"),
    ("✅","Build log — successful","Zero errors"),
    ("✅","Live URL generated","onrender.com active"),
    ("✅","Login page verified on live URL","Screenshot: Fig 5"),
    ("✅","Admin dashboard tested","Screenshot: Fig 6"),
    ("✅","Student dashboard tested","Screenshot: Fig 7"),
    ("✅","60% gap rule validated live","Gap detection working"),
    ("✅","Faculty demo ready","Project Complete"),
]:
    row=chk.add_row()
    for cell,(txt,al) in zip(row.cells,[(s,WD_ALIGN_PARAGRAPH.CENTER),(t,WD_ALIGN_PARAGRAPH.LEFT),(o,WD_ALIGN_PARAGRAPH.LEFT)]):
        p=cell.paragraphs[0]; p.alignment=al
        p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        sf(p.add_run(txt),10.5)

cp(doc); hrule(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 7 — TECH STACK
# ══════════════════════════════════════════════════════════════
cp(doc,"7. Production Technology Stack",14,True,(0,60,120),sb=6,sa=6)
tt=doc.add_table(rows=1,cols=3); tt.style='Table Grid'; tt.alignment=WD_TABLE_ALIGNMENT.CENTER
for cell,txt in zip(tt.rows[0].cells,["Layer","Technology","Role"]):
    set_bg(cell,"003366"); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt),11,True,(255,255,255))
for l,t,r in [
    ("Cloud","Render.com (Free Tier)","Production cloud hosting"),
    ("WSGI","Gunicorn","Python application server"),
    ("Backend","Flask (Python 3.11)","Web application framework"),
    ("Database","MongoDB Atlas","Cloud NoSQL database + GridFS"),
    ("Auth","Google OAuth (Authlib)","Secure Single Sign-On"),
    ("AI","Google Gemini AI","Skill gap analysis engine"),
    ("Videos","YouTube Data API v3","Upskilling recommendations"),
    ("CI/CD","GitHub + Render Auto-Deploy","Auto deploy on git push"),
]:
    row=tt.add_row()
    for cell,txt in zip(row.cells,[l,t,r]):
        p=cell.paragraphs[0]; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        sf(p.add_run(txt),10.5)

cp(doc); cp(doc)
cp(doc,"GapLens Platform  •  Week 12 — Deployment & Final Review  •  Team B-G12  •  SVIIT  •  2026",9,color=(150,150,150),align=WD_ALIGN_PARAGRAPH.CENTER)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"Week_12_Render_Deployment.docx")
doc.save(out)
print(f"[OK] Saved: {out}")
