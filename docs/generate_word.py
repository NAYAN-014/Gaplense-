from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BRAIN = r"C:\Users\Nayan Patidar\.gemini\antigravity\brain\5562ef76-3b16-4013-b3d7-41c9342a8ac7"
IMGS = {
    "timeline":   os.path.join(BRAIN, "project_timeline_1777810416421.png"),
    "arch":       os.path.join(BRAIN, "system_architecture_1777810430742.png"),
    "cicd":       os.path.join(BRAIN, "cicd_pipeline_1777810578079.png"),
    "usecase":    os.path.join(BRAIN, "use_case_diagram_1777810644552.png"),
    "class":      os.path.join(BRAIN, "uml_class_diagram_1777811231157.png"),
    "sequence":   os.path.join(BRAIN, "sequence_diagram_1777811291450.png"),
}

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

def sf(run, sz=11, bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def cp(doc, text="", sz=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        sf(r, sz, bold, color, italic)
    return p

def add_img(doc, key, caption, w=Inches(5.5)):
    path = IMGS.get(key, "")
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.add_run().add_picture(path, width=w)
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        sf(r, 9, italic=True, color=(100,100,100))
        c.paragraph_format.space_after = Pt(10)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_para(cell, parts, sz=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_b=1, space_a=1, indent=0):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_b)
    p.paragraph_format.space_after  = Pt(space_a)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    for txt, bold in parts:
        r = p.add_run(txt)
        sf(r, sz, bold)
    return p

# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
cp(doc, "SHRI VAISHNAV VIDYAPEETH VISHWAVIDYALAYA", 16, True, (0,60,120), WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
cp(doc, "Shri Vaishnav Institute of Information Technology", 13, True, (0,60,120), WD_ALIGN_PARAGRAPH.CENTER)
cp(doc)
cp(doc, "Ref. No.: SVIIT/2026/01/80", 11, False, align=WD_ALIGN_PARAGRAPH.LEFT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Date: 13/01/2026")
sf(r, 11)

cp(doc)
cp(doc, "NOTICE", 14, True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
for side in ['top','bottom']:
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'8')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'000000')
    pBdr.append(b)
pPr.append(pBdr)
r = p.add_run("Minor Project Documentation — GapLens Skill Diagnostic Platform")
sf(r, 13, True)

cp(doc)
add_img(doc, "timeline", "Figure 1: 12-Week Project Timeline", w=Inches(6.0))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1: PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════
cp(doc, "1. Project Overview", 14, True, (0,60,120), space_before=6, space_after=4)
cp(doc, "GapLens is an AI-powered skill diagnostic and career gap analysis platform. It identifies skill gaps between a student's current abilities and industry requirements, then recommends personalized upskilling resources using YouTube API and Google Gemini AI.", 11, space_after=4)

p = doc.add_paragraph()
sf(p.add_run("Project Title: "), 11, True)
sf(p.add_run("GapLens — Skill Diagnostic & Employment Platform"), 11)

p = doc.add_paragraph()
sf(p.add_run("Technology Stack: "), 11, True)
sf(p.add_run("Python 3.14 · Flask · MongoDB · GridFS · Google OAuth · Gemini AI · YouTube API · Tailwind CSS"), 11)

p = doc.add_paragraph()
sf(p.add_run("Roles: "), 11, True)
sf(p.add_run("Student · Recruiter · Admin"), 11)

p = doc.add_paragraph()
sf(p.add_run("Version Control: "), 11, True)
sf(p.add_run("GitHub — github.com/NAYAN-014/Gaplense-"), 11)

cp(doc)
add_img(doc, "arch", "Figure 2: System Architecture — Flask Blueprints, MongoDB, External APIs", w=Inches(5.8))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 2: UML DIAGRAMS
# ══════════════════════════════════════════════════════════════════
cp(doc, "2. UML Diagrams", 14, True, (0,60,120), space_before=6, space_after=6)

cp(doc, "2.1  Use Case Diagram", 12, True, (0,80,150), space_after=4)
cp(doc, "Defines system boundaries and interactions between actors (Student, Recruiter, Admin) and system functionalities.", 11, space_after=4)
add_img(doc, "usecase", "Figure 3: Use Case Diagram — All Roles", w=Inches(5.8))

cp(doc)
cp(doc, "2.2  Class Diagram (Domain Model)", 12, True, (0,80,150), space_after=4)
cp(doc, "Shows all domain entities, their attributes, methods, and relationships. Key classes: User, Assessment, Result, Job, Resume, CommunityPost.", 11, space_after=4)
add_img(doc, "class", "Figure 4: Class Diagram — GapLens Domain Model", w=Inches(5.8))

cp(doc)
cp(doc, "2.3  Sequence Diagram — Diagnostic Test Flow", 12, True, (0,80,150), space_after=4)
cp(doc, "Shows how the Student, Browser, Flask Backend, MongoDB, and Gemini AI interact during a diagnostic test session.", 11, space_after=4)
add_img(doc, "sequence", "Figure 5: Sequence Diagram — Diagnostic Test Flow", w=Inches(5.8))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 3: PROJECT TIMELINE TABLE
# ══════════════════════════════════════════════════════════════════
cp(doc, "3. Minor Project Timeline", 14, True, (0,60,120), space_before=6, space_after=6)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.columns[0].width = Inches(4.8)
tbl.columns[1].width = Inches(1.8)

# Header row
hdr = tbl.rows[0]
set_cell_bg(hdr.cells[0], "003366")
set_cell_bg(hdr.cells[1], "003366")
for cell, txt in zip(hdr.cells, ["Phases", "Date"]):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    sf(r, 12, True, (255,255,255))

def add_row(txt_parts_left, date_txt, bg="FFFFFF", date_bg="F0F7FF"):
    row = tbl.add_row()
    row.cells[0].width = Inches(4.8)
    row.cells[1].width = Inches(1.8)
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], date_bg)
    lc = row.cells[0]
    for parts, indent, sz, sb, sa in txt_parts_left:
        p = lc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        for txt, bold in parts:
            r = p.add_run(txt)
            sf(r, sz, bold)
    rc = row.cells[1]
    p = rc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    r = p.add_run(date_txt)
    sf(r, 9.5, False, (0,60,120))
    return row

# ── PHASE 1 ──────────────────────────────────────────────────────
add_row([
    ([("Phase 1: Planning & Requirement Analysis (Weeks 1–2)", True)], 0, 11.5, 4, 2),
    ([("Focus: ", True), ("Defining the scope, setting up infrastructure, and identifying functional requirements.", False)], 0, 10, 2, 4),
], "", "D6E8FF", "D6E8FF")

add_row([
    ([("●  Week 1: Project Initiation", True)], 0.1, 11, 3, 1),
    ([("○  Form teams and define the problem statement.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("Tooling: ", True), ("Set up ", False), ("Jira", True), (" boards (Scrum/Kanban), create Backlog, initialize ", False), ("Git", True), (" repositories.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("Deliverable: ", True), ("Project Vision document and initial Product Backlog.", False)], 0.3, 10, 1, 3),
], "19–24 January\n2026")

add_row([
    ([("●  Week 2: Functional Modeling", True)], 0.1, 11, 3, 1),
    ([("○  Identify Actors and Use Cases.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("UML Deliverable: Use Case Diagram", True), (" — To define system boundaries.", False)], 0.3, 10, 1, 1),
    ([("     ■  ", False), ("Use Case Descriptions: ", True), ("Detailed text for each primary flow.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("Tooling: ", True), ("Map Use Cases to Jira User Stories.", False)], 0.3, 10, 1, 3),
], "26–31 January\n2026")

# ── PHASE 2 ──────────────────────────────────────────────────────
add_row([
    ([("Phase 2: Object-Oriented Analysis (Weeks 3–4)", True)], 0, 11.5, 4, 2),
    ([("Focus: ", True), ('Understanding the "What" of the system without focusing on implementation.', False)], 0, 10, 2, 4),
], "", "D6E8FF", "D6E8FF")

add_row([
    ([("●  Week 3: Static Analysis", True)], 0.1, 11, 3, 1),
    ([("○  Identify domain entities and their relationships.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("UML Deliverables:", True)], 0.3, 10, 1, 1),
    ([("     ■  ", False), ("Object Diagram: ", True), ("Model specific instances and data snapshots.", False)], 0.3, 10, 1, 1),
    ([("     ■  ", False), ("Initial Class Diagram (Domain Model): ", True), ("Attributes and relationships (1:1, 1:N, M:N).", False)], 0.3, 10, 1, 3),
], "02–07 February\n2026")

add_row([
    ([("●  Week 4: Dynamic Analysis", True)], 0.1, 11, 3, 1),
    ([("○  Model how the system reacts to events.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("UML Deliverable — Activity Diagram: ", True), ("Model complex business logic and parallel workflows.", False)], 0.3, 10, 1, 3),
], "09–14 February\n2026")

# ── PHASE 3 ──────────────────────────────────────────────────────
add_row([
    ([("Phase 3: Object-Oriented Design (Weeks 5–7)", True)], 0, 11.5, 4, 2),
    ([("Focus: ", True), ('Defining the "How" — architecture, patterns, and component interactions.', False)], 0, 10, 2, 4),
], "", "D6E8FF", "D6E8FF")

add_row([
    ([("●  Week 5: Interaction Modeling", True)], 0.1, 11, 3, 1),
    ([("○  ", False), ("UML Deliverable — Sequence Diagram: ", True), ("Detailed logic showing how objects communicate over time to fulfill a use case.", False)], 0.3, 10, 1, 3),
], "16–21 February\n2026")

add_row([
    ([("●  Week 6: Detailed Class Design & Packaging", True)], 0.1, 11, 3, 1),
    ([("○  Apply SOLID principles and Design Patterns.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("UML Deliverables:", True)], 0.3, 10, 1, 1),
    ([("     ■  ", False), ("Detailed Class Diagram: ", True), ("Including methods, visibility, and interfaces.", False)], 0.3, 10, 1, 1),
    ([("     ■  ", False), ("Package Diagram: ", True), ("Organizing classes into logical modules/subsystems.", False)], 0.3, 10, 1, 3),
], "23–28 February\n2026")

add_row([
    ([("●  Week 7: System Architecture", True)], 0.1, 11, 3, 1),
    ([("○  ", False), ("UML Deliverables:", True)], 0.3, 10, 1, 1),
    ([("     ■  ", False), ("Component Diagram: ", True), ("Visualizing physical modules (API, Database, UI).", False)], 0.3, 10, 1, 1),
    ([("     ■  ", False), ("Deployment Diagram: ", True), ("Mapping software components to hardware/nodes (Cloud, Servers).", False)], 0.3, 10, 1, 3),
], "02–07 March\n2026")

# ── PHASE 4 ──────────────────────────────────────────────────────
add_row([
    ([("Phase 4: Implementation & CI/CD Integration (Weeks 8–10)", True)], 0, 11.5, 4, 2),
    ([("Focus: ", True), ("Coding, version control, and automated pipelines.", False)], 0, 10, 2, 4),
], "", "D6E8FF", "D6E8FF")

add_row([
    ([("●  Week 8: Sprint 1 — Core Features", True)], 0.1, 11, 3, 1),
    ([("○  Develop core logic based on the Class and Sequence diagrams.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("Git Strategy: ", True), ("Use branching (Feature branches → Develop → Main).", False)], 0.3, 10, 1, 3),
], "09–14 March\n2026")

add_row([
    ([("●  Week 9: CI/CD Pipeline Setup", True)], 0.1, 11, 3, 1),
    ([("○  Configure ", False), ("GitHub Actions", True), (" or ", False), ("GitLab CI.", True)], 0.3, 10, 1, 1),
    ([("○  Automate build triggers on every git push.", False)], 0.3, 10, 1, 3),
], "16–21 March\n2026")

add_row([
    ([("●  Week 10: Sprint 2 — Integration", True)], 0.1, 11, 3, 1),
    ([("○  Develop UI and connect to the database.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("Jira: ", True), ('Move stories to "In Progress" and "Testing."', False)], 0.3, 10, 1, 3),
], "23–28 March\n2026")

# ── PHASE 5 ──────────────────────────────────────────────────────
add_row([
    ([("Phase 5: Testing & Deployment (Weeks 11–12)", True)], 0, 11.5, 4, 2),
    ([("Focus: ", True), ("Quality assurance and final delivery.", False)], 0, 10, 2, 4),
], "", "D6E8FF", "D6E8FF")

add_row([
    ([("●  Week 11: Testing Phase", True)], 0.1, 11, 3, 1),
    ([("○  ", False), ("Unit Testing: ", True), ("(JUnit, PyTest) ensuring methods work as intended.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("Integration Testing: ", True), ("Testing the flow between components.", False)], 0.3, 10, 1, 1),
    ([("○  ", False), ("Tooling: ", True), ("Integrate test reports into the CI/CD pipeline.", False)], 0.3, 10, 1, 3),
], "30 March–\n04 April 2026")

add_row([
    ([("●  Week 12: Deployment & Final Review", True)], 0.1, 11, 3, 1),
    ([("○  Finalize the ", False), ("Deployment Diagram", True), (" to match the production environment.", False)], 0.3, 10, 1, 1),
    ([("○  Deploy to staging/production (Heroku, AWS, or Azure).", False)], 0.3, 10, 1, 1),
    ([("○  Project Demo and handover.", False)], 0.3, 10, 1, 3),
], "06–11 April\n2026")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 4: CI/CD PIPELINE
# ══════════════════════════════════════════════════════════════════
cp(doc, "4. CI/CD Pipeline", 14, True, (0,60,120), space_before=6, space_after=4)
cp(doc, "The project uses GitHub Actions for continuous integration. On every git push to the main branch, the pipeline automatically runs tests (PyTest), validates imports, and deploys to the staging environment.", 11, space_after=6)
add_img(doc, "cicd", "Figure 6: CI/CD Pipeline — Automated Build, Test & Deploy Flow", w=Inches(5.8))

cp(doc)
cp(doc, "4.1  Git Branching Strategy", 12, True, (0,80,150), space_after=4)
for line in ["● main — Production-ready code only", "● develop — Integration branch for tested features", "● feature/* — Individual feature development branches"]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    sf(p.add_run(line), 11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 5: PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════
cp(doc, "5. Project Folder Structure", 14, True, (0,60,120), space_before=6, space_after=4)
struct = [
    ("GapLens/", True), ("├── run.py              ← Entry point", False),
    ("├── main/app.py         ← Flask app factory", False),
    ("├── other/              ← config.py, extensions.py", False),
    ("├── routes/             ← auth, student, recruiter, admin, diagnostic, resume", False),
    ("├── core/               ← utils.py, seed_questions.py", False),
    ("├── static/css/         ← style.css", False),
    ("├── static/img/         ← logo.png", False),
    ("├── templates/          ← base.html, auth/, student/, recruiter/, admin/, test/", False),
    ("└── docs/               ← Project_Documentation.md, GapLens_Project_Plan.docx", False),
]
for txt, bold in struct:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(txt)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.bold = bold
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)

cp(doc)
cp(doc, "6. Technology Stack", 14, True, (0,60,120), space_before=10, space_after=6)

tech_tbl = doc.add_table(rows=1, cols=3)
tech_tbl.style = 'Table Grid'
tech_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, txt in zip(tech_tbl.rows[0].cells, ["Layer", "Technology", "Purpose"]):
    set_cell_bg(cell, "003366")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run(txt), 11, True, (255,255,255))

rows = [
    ("Backend", "Flask (Python 3.14)", "Web framework with Blueprint modular routing"),
    ("Database", "MongoDB + GridFS", "NoSQL storage for users, results, resumes"),
    ("Auth", "Google OAuth (Authlib)", "Single Sign-On + local session management"),
    ("AI / ML", "Google Gemini AI", "Skill gap analysis and AI insights generation"),
    ("Recommendations", "YouTube Data API v3", "Fetch curated learning videos per skill gap"),
    ("Frontend", "Tailwind CSS + Jinja2", "Responsive UI with dark/light mode support"),
    ("Deployment", "GitHub + Heroku / AWS", "Version control and cloud deployment"),
]
for layer, tech, purpose in rows:
    row = tech_tbl.add_row()
    for cell, txt in zip(row.cells, [layer, tech, purpose]):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        sf(p.add_run(txt), 10.5)

# Footer
cp(doc)
cp(doc)
cp(doc, "GapLens Platform  •  Minor Project Documentation  •  SVIIT  •  2026", 9,
   color=(150,150,150), align=WD_ALIGN_PARAGRAPH.CENTER)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "GapLens_Project_Plan.docx")
doc.save(out)
print(f"[OK] Saved: {out}")
